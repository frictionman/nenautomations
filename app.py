
# Developed by Nen AI

from fetch_image import fetch_photo
from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
from docx import Document
from docx.shared import Inches
import io
from PIL import Image

# Function to load the Llama 2 model
def load_llm(max_tokens, prompt_template):
    # Initialize the Llama 2 model with given parameters
    llm = CTransformers(
        model = "llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens = max_tokens,
        temperature = 0.7
    )
    
    # Set up a "chain" for processing prompts
    llm_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(prompt_template)
    )
    return llm_chain

# Function to create a Word document with user input, generated article, and image
def create_word_docx(user_input, paragraph, image_input):
    doc = Document()
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))
    return doc

# Streamlit UI setup
def main():
    st.title("Article Generator using Llama 2 by Nen AI")
    
    # Take input from the user for article topic and image topic
    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
    image_input = st.text_input("Please enter the topic for the image you want to fetch!")
    
    if user_input and image_input:
        # Display generated content and fetched image side by side
        col1, col2, col3 = st.columns([1,2,1])
        
        # Generate article using Llama 2
        with col1:
            prompt_template = f"You are a digital marketing and SEO expert and your task is to write article so write an article on the given topic: {user_input}. The article should be under 800 words."
            llm_chain = load_llm(max_tokens=800, prompt_template=prompt_template)
            result = llm_chain(user_input)
            if result:
                st.write(result)
            else:
                st.error("Your article couldn't be generated!")
        
        # Fetch image using the fetch_photo function
        with col2:
            image_url = fetch_photo(image_input)
            if image_url:
                st.image(image_url)
            else:
                st.write("No image found for the given topic.")
        
        # Allow user to download the final article with the image
        with col3:
            doc = create_word_docx(user_input, result['text'], Image.open(requests.get(image_url, stream=True).raw))
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

if __name__ == "__main__":
    main()
